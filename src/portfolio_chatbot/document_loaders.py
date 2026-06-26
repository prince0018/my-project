from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


MetadataValue = str | int | float | bool
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class SourceDocument:
    text: str
    metadata: dict[str, MetadataValue]


def iter_document_files(docs_path: Path) -> list[Path]:
    if not docs_path.exists():
        return []

    return [
        path
        for path in sorted(docs_path.rglob("*"))
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
        and not _is_docs_readme(path, docs_path)
    ]


def load_documents_from_directory(docs_path: Path) -> list[SourceDocument]:
    documents: list[SourceDocument] = []

    for path in iter_document_files(docs_path):
        documents.extend(load_document(path, docs_path))

    return documents


def load_document(path: Path, docs_path: Path) -> list[SourceDocument]:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path, docs_path)
    if suffix == ".docx":
        return _load_docx(path, docs_path)
    if suffix in {".txt", ".md"}:
        return _load_text_file(path, docs_path)

    return []


def _load_pdf(path: Path, docs_path: Path) -> list[SourceDocument]:
    reader = PdfReader(str(path))
    documents: list[SourceDocument] = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue

        documents.append(
            SourceDocument(
                text=text,
                metadata={
                    **_base_metadata(path, docs_path),
                    "page": page_number,
                },
            )
        )

    return documents


def _load_docx(path: Path, docs_path: Path) -> list[SourceDocument]:
    document = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        return []

    return [
        SourceDocument(
            text=text,
            metadata=_base_metadata(path, docs_path),
        )
    ]


def _load_text_file(path: Path, docs_path: Path) -> list[SourceDocument]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []

    return [
        SourceDocument(
            text=text,
            metadata=_base_metadata(path, docs_path),
        )
    ]


def _base_metadata(path: Path, docs_path: Path) -> dict[str, MetadataValue]:
    return {
        "source": str(path.relative_to(docs_path)),
        "file_name": path.name,
        "file_type": path.suffix.lower().lstrip("."),
    }


def _is_docs_readme(path: Path, docs_path: Path) -> bool:
    return path.resolve() == (docs_path / "README.md").resolve()
